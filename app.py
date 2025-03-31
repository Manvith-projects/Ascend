from flask import Flask, render_template, request, redirect, url_for, session, flash,jsonify, send_from_directory
import os
import requests
import textract
import fitz  # PyMuPDF
from dotenv import load_dotenv
import json,io
from werkzeug.utils import secure_filename
import google.generativeai as genai
import joblib
import PyPDF2
from docx import Document
import numpy as np
from gensim.models import Word2Vec
import tensorflow as tf
import numpy as np
import pickle
import json
import pandas as pd
import random
import nltk
from nltk.stem import WordNetLemmatizer
from flask_cors import CORS
from PyPDF2 import PdfReader
import subprocess

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")

# ----------------------------------------------------------------------------login req----------------------------------------------------------------

#------------------------------------------------------------------------------Routes----------------------------------------------------------------
@app.route('/')
def index():
    return render_template('index.html')  

@app.route('/login')
def login():
    return render_template('./auth/login.html')

@app.route('/signup')
def signup():
    return render_template('./auth/signup.html')
  
@app.route('/logout')
def logout():
    session.clear()  
    return redirect(url_for('login'))  

@app.route('/landing')

def landing():
    return render_template('landing.html')

@app.route('/interview')

def interview():
    return render_template('interview.html')

@app.route('/still_in_progress')

def still_in_progress():
    return render_template('still_in_progress.html')


@app.route('/forgotpassword')

def forgotpassword():
    return render_template('./auth/forgotpassword.html')

@app.route('/get-started')

def get_started():
    return render_template('get_started.html')

@app.route('/careerconfusion')

def careerconfusion():
    return render_template('careerconfusion.html')

@app.route('/learn')

def learn():
    return render_template('learn.html')

@app.route('/chatbot')

def chatbot():
    return render_template('chatbot.html')

@app.route('/JobSearch')

def JobSearch():
    return render_template('JobSearch.html')

@app.route('/startup')

def startup():
    return render_template('startup.html')

@app.route('/trend')

def trend():
    return render_template('trend.html')

@app.route('/PublicSector')

def PublicSector():
    return render_template('PublicSector.html')

@app.route('/mentor')

def mentor():
    return render_template('mentor.html')

@app.route('/Private')

def Private():
    return render_template('Private.html')

@app.route('/nong')

def nong():
    return render_template('nong.html')

@app.route('/counsillingBot')

def counsillingBot():
    return render_template('counsillingBot.html')

@app.route('/resumepredict')

def resumepredict():
    return render_template('resumepredict.html')

@app.route('/jobs')

def jobs():
    return render_template('jobs.html') 

@app.route('/AfterSchool')

def AfterSchool():
    return render_template('afterSchool.html') 

@app.route('/ai.html')

def ai():
  return render_template('./roles/ai.html')

@app.route('/SysArc.html')

def sys_arc():
  return render_template('./roles/SysArc.html')

@app.route('/Projman.html')

def proj_man():
  return render_template('./roles/Projman.html')

@app.route('/ML.html')

def ml():
  return render_template('./roles/ML.html')

@app.route('/Data_Eng.html')

def data_eng():
  return render_template('./roles/Data_Eng.html')

@app.route('/ITPROJMAN.html')

def it_proj_man():
  return render_template('./roles/ITPROJMAN.html')

@app.route('/Bussiness.html')

def business():
  return render_template('./roles/Bussiness.html')

@app.route('/NetAdm.html')

def net_adm():
  return render_template('./roles/NetAdm.html')

@app.route('/Quality.html')

def quality():
  return render_template('./roles/Quality.html')

@app.route('/CSA.html')

def csa():
  return render_template('./roles/CSA.html')

@app.route('/BLKCH.html')

def blkch():
  return render_template('./roles/BLKCH.html')

@app.route('/Cyber.html')

def cyber():
  return render_template('./roles/Cyber.html')

@app.route('/FullStack.html')

def full_stack():
  return render_template('./roles/FullStack.html')

@app.route('/IOT.html')

def iot():
  return render_template('./roles/IOT.html')

@app.route('/NLP.html')

def nlp():
  return render_template('./roles/NLP.html')

@app.route('/MObile.html')

def mobile():
  return render_template('./roles/MObile.html')

@app.route('/Game.html')

def game():
  return render_template('./roles/Game.html')

@app.route('/UX.html')

def ux():
  return render_template('./roles/UX.html')

@app.route('/Quant.html')

def quant():
  return render_template('./roles/Quant.html')

@app.route('/Robotics.html')

def robotics():
  return render_template('./roles/Robotics.html')

@app.route('/ITsec.html')

def it_sec():
  return render_template('./roles/ITsec.html')

@app.route('/HR.html')

def hr():
  return render_template('./roles/HR.html')

@app.route('/CTO.html')

def cto():
  return render_template('./roles/CTO.html')

@app.route('/BigData.html')

def big_data():
  return render_template('./roles/BigData.html')

@app.route('/TechSup.html')


def tech_sup():
  return render_template('./roles/TechSup.html')

@app.route('/Prompt.html')

def prompt():
  return render_template('./roles/Prompt.html')

@app.route('/Nurse.html')

def nurse():
  return render_template('./roles/Nurse.html')

@app.route('/mediLab.html')

def medi_lab():
  return render_template('./roles/mediLab.html')

@app.route('/healthAdmin.html')

def health_admin():
  return render_template('./roles/healthAdmin.html')

@app.route('/MedicalAsst.html')

def medical_asst():
  return render_template('./roles/MedicalAsst.html')

@app.route('/FinAnalyst.html')

def fin_analyst():
  return render_template('./roles/FinAnalyst.html')

@app.route('/CA.html')

def ca():
  return render_template('./roles/CA.html')

@app.route('/FInPlan.html')

def fin_plan():
  return render_template('./roles/FInPlan.html')

@app.route('/InvesBnk.html')

def inves_bnk():
  return render_template('./roles/InvesBnk.html')

@app.route('/Risk.html')

def risk():
  return render_template('./roles/Risk.html')

@app.route('/Commercial.html')

def commercial():
  return render_template('./roles/Commercial.html')

@app.route('/Loan.html')

def loan():
  return render_template('./roles/Loan.html')

@app.route('/Wealth.html')

def wealth():
  return render_template('./roles/Wealth.html')

@app.route('/Relation.html')

def relation():
  return render_template('./roles/Relation.html')

@app.route('/Recovery.html')

def recovery():
  return render_template('./roles/Recovery.html')

@app.route('/BussinessDev.html')

def business_dev():
  return render_template('./roles/BussinessDev.html')

@app.route('/DigitalMarkrtingSpecilist.html')

def digital_marketing_specialist():
  return render_template('./roles/DigitalMarkrtingSpecilist.html')

@app.route('/SalesAcountExecutive.html')

def sales_account_executive():
  return render_template('./roles/SalesAcountExecutive.html')

@app.route('/MarketingManager.html')

def marketing_manager():
  return render_template('./roles/MarketingManager.html')

@app.route('/DataScientist.html')

def data_scientist():
  return render_template('./roles/DataScientist.html')

@app.route('/ResearchScientist.html')

def research_scientist():
  return render_template('./roles/ResearchScientist.html')

@app.route('/ClinicalResearcAsso.html')

def clinical_research_associate():
  return render_template('./roles/ClinicalResearcAsso.html')

@app.route('/InnovationManager.html')

def innovation_manager():
  return render_template('./roles/InnovationManager.html')

@app.route('/Teacher.html')

def teacher():
  return render_template('./roles/Teacher.html')

@app.route('/CarriculamDes.html')

def curricula_designer():
  return render_template('./roles/CarriculamDes.html')

@app.route('/SpecialEducation.html')

def special_education():
  return render_template('./roles/SpecialEducation.html')

@app.route('/councilor.html')

def councilor():
  return render_template('./roles/councilor.html')

@app.route('/ProductionManager.html')

def production_manager():
  return render_template('./roles/ProductionManager.html')

@app.route('/QualityControl.html')

def quality_control():
  return render_template('./roles/QualityControl.html')

@app.route('/IndusEngineer.html')

def industrial_engineer():
  return render_template('./roles/IndusEngineer.html')

@app.route('/PlantManager.html')

def plant_manager():
  return render_template('./roles/PlantManager.html')

@app.route('/SupplyChainCoor.html')


def supply_chain_coordinator():
  return render_template('./roles/SupplyChainCoor.html')

@app.route('/LogisManager.html')

def logistics_manager():
  return render_template('./roles/LogisManager.html')

@app.route('/WareHouseManager.html')

def warehouse_manager():
  return render_template('./roles/WareHouseManager.html')

@app.route('/Procurespelist.html')

def procurement_specialist():
  return render_template('./roles/Procurespelist.html')

@app.route('/ConstructionManager.html')

def construction_manager():
  return render_template('./roles/ConstructionManager.html')

@app.route('/Civil.html')

def civil():
  return render_template('./roles/Civil.html')

@app.route('/Architect.html')
def architect():
  return render_template('./roles/Architect.html')

@app.route('/Survey.html')
def survey():
  return render_template('./roles/Survey.html')

@app.route('/GeneralManager.html')
def general_manager():
  return render_template('./roles/GeneralManager.html')

@app.route('/HouseKeeping.html')
def housekeeping():
  return render_template('./roles/HouseKeeping.html')

@app.route('/Event.html')
def event():
  return render_template('./roles/Event.html')

@app.route('/Food&Bevarages.html')
def food_beverages():
  return render_template('./roles/Food&Bevarages.html')

@app.route('/profile')
def profile():
  return render_template('Profile.html')

@app.route('/own.html')
def own():
  return render_template('own.html')

@app.route('/news.html')
def news():
  return render_template('news.html')
# @app.route('/news.html')
# def news():
#   return render_template('news.html')
@app.route('/joinus.html')
def joinus():
  return render_template('join_us.html')

@app.route('/privacypolicy-2.html')
def privacypolicy():
  return render_template('privacy-policy-2.html')

@app.route('/Terms.html')
def Terms():
  return render_template('Terms of Service.html')

@app.route('/contact.html')
def contact():
  return render_template('ContactUs.html')


@app.route('/About_Us.html')
def aboutus():
  return render_template('About_Us.html')

@app.route('/connect.html')
def connect():
  return render_template('connect.html')


@app.route('/explore')
def explore():
    
    return render_template('explore.html')  

#-----------------------------------------------------------------------------------chatbot----------------------------------------------------------------------


def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text

pdf_text = extract_text_from_pdf("static/dataset/About_Carer_Horizon.pdf")



# ✅ Configure Google Gemini API
genai.configure(api_key="AIzaSyAyD4hCI7hfXnZ1RX9XfalqJNFwHeV9RFE")
model = genai.GenerativeModel("gemini-1.5-flash")

@app.route("/chatbotw", methods=["POST"])
def chatbotw():
    user_input = request.json.get("message", "")

    if not user_input:
        return jsonify({"error": "No message provided"}), 400

    prompt = f"""
    You are a chatbot for Career Horizon. Use the following website details to assist users:
    {pdf_text}

    User: {user_input}
    Chatbot:
    """
    
    try:
        response = model.generate_content(prompt)
        bot_reply = response.text.strip()
    except Exception as e:
        bot_reply = "Sorry, I couldn't fetch an answer."

    return jsonify({"response": bot_reply})
    







#-------------------------------------------------------------------------------------blog routes---------------------------------------------------
@app.route('/blogs')
def blogs():
    
    return render_template('./blog/blog.html') 

 


#----------------------------------------------------------------job search---------------------------------------------
RAPIDAPI_KEY = "9a6fafe68cmshca09d2d8e9dbb59p14e3d5jsn2f64c416e74b"
RAPIDAPI_HOST = "jsearch.p.rapidapi.com"

@app.route('/search-jobs', methods=['GET'])
def search_jobs():
    title = request.args.get('title')
    location = request.args.get('location')
    page = request.args.get('page', 1)

    # Corrected API URL
    api_url = f"https://jsearch.p.rapidapi.com/search?query={title}+in+{location}&page={page}"

    headers = {
        "X-RapidAPI-Key": RAPIDAPI_KEY,
        "X-RapidAPI-Host": RAPIDAPI_HOST
    }

    # 🛑 You were using response.json() before making the request! 
    response = requests.get(api_url, headers=headers)
    
    # Convert to JSON
    response_data = response.json()

    # Print locations for debugging
    if "data" in response_data:
        for job in response_data["data"]:
            print(job.get("location", "No Location Found"))

    return jsonify(response_data)  # Send data to frontend

#-------------------------------------------------------------------------------------resume prediction---------------------------------------------------
# Configure Gemini API
genai.configure(api_key="AIzaSyBSwVAIWGui-QAxEzufmcpVH-RvWjrNmqc") # Replace with your Gemini API key
model = genai.GenerativeModel('gemini-1.5-pro')

# RapidAPI JSearch API Configuration
JSEARCH_API_KEY = "9a6fafe68cmshca09d2d8e9dbb59p14e3d5jsn2f64c416e74b" # Replace with your Rapid JSearch API key
JSEARCH_API_HOST = "jsearch.p.rapidapi.com"

def extract_text_from_pdf(file_stream):
    """Extracts text from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
            print(text)
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {e}"

def extract_text_from_docx(file_stream):
    """Extracts text from a Word document."""
    try:
        doc = Document(file_stream)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        return f"Error extracting text from Word document: {e}"

def analyze_resume_with_gemini(resume_text):
    """Analyzes resume text with Gemini API and extracts top 3 skills."""
    prompt = f"Analyze the resume text below and identify the top 3 most relevant job roles based on the candidate's skills, experience, and industry trends. Return only the job titles as a comma-separated list, without any additional text:\n\n{resume_text}"

    try:
        print(resume_text)
        print("Sending request to Gemini API...")
        response = model.generate_content(prompt)
        print(response)
        print("Gemini API response received.")
        skills_string = response.text.strip()
        skills = [skill.strip() for skill in skills_string.split(',')]
        
        print(skills)
        return skills  # Return all extracted skills
    except Exception as e:
        print(f"Gemini API Error: {e}")
        return [f"Error analyzing resume: {e}"]

def search_jobs_with_jsearch(keywords):
    """Searches for jobs with Rapid JSearch API."""
    url = "https://jsearch.p.rapidapi.com/search"
    querystring = {"query": ", ".join(keywords), "num_pages": "1"}
    headers = {
        "X-RapidAPI-Key": JSEARCH_API_KEY,
        "X-RapidAPI-Host": JSEARCH_API_HOST,
    }
    try:
        response = requests.get(url, headers=headers, params=querystring)
        response.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)
        print(response.json().get('data', []))
        return response.json().get('data', [])
    except requests.exceptions.RequestException as e:
        return [f"Error fetching jobs: {e}"]

@app.route('/resumeprediction', methods=['GET', 'POST'])
def resumeprediction():
    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if file:
            file_stream = io.BytesIO(file.read())
            if file.filename.lower().endswith('.pdf'):
                resume_text = extract_text_from_pdf(file_stream)
            elif file.filename.lower().endswith('.docx'):
                resume_text = extract_text_from_docx(file_stream)
            else:
                return jsonify({"error": "Invalid file format. Please upload a PDF or DOCX file."}), 400

            print("Uploading")
            skills = analyze_resume_with_gemini(resume_text)
            jobs = search_jobs_with_jsearch(skills)
            print(jsonify({"jobs": jobs, "skills": skills}))
            return jsonify({"jobs": jobs, "skills": skills})

    return render_template('resumepredict.html')

    
#-------------------------------------------------------------------------------------ML model---------------------------------------------------
# # Load models
# rf_model = joblib.load("rf_model.joblib", mmap_mode='r')
# word2vec_model = Word2Vec.load("word22vec_model.model")

# # ✅ Function to get word embeddings
# def get_embeddings(doc):
#     vectors = [word2vec_model.wv[word] for word in doc if word in word2vec_model.wv]
#     return sum(vectors) if vectors else np.zeros(word2vec_model.vector_size)

# # ✅ Function to preprocess input
# def preprocess_input(input_text):
#     tokenized_input = input_text.split()
#     input_embeddings = get_embeddings(tokenized_input)
#     return input_embeddings

# # ✅ Function to predict role
# def predict_role(input_text):
#     input_embeddings = preprocess_input(input_text)
#     predicted_role = rf_model.predict([input_embeddings])
#     return predicted_role[0]

# # ✅ Route to handle job role prediction
# @app.route("/predict", methods=["GET", "POST"])
# def jobml():
#     if request.method == "POST":
#         # Check if request is JSON (AJAX) or Form Submission
#         if request.is_json:
#             data = request.get_json()
#             industry_type = data.get("industry_type", "").strip()
#             department = data.get("department", "").strip()
#             role_category = data.get("role_category", "").strip()
#             key_skills = data.get("key_skills", "").strip()
#         else:
#             industry_type = request.form.get("industry_type", "").strip()
#             department = request.form.get("department", "").strip()
#             role_category = request.form.get("role_category", "").strip()
#             key_skills = request.form.get("key_skills", "").strip()

#         # Combine inputs
#         user_input = f"{industry_type} {department} {role_category} {key_skills}"

#         # Predict role
#         predicted_role = predict_role(user_input)
#         print("Predicted Role:", predicted_role)

#         # Return response based on request type
#         if request.is_json:
#             return jsonify({"predicted_role": predicted_role})
#         else:
#             return render_template("JobSearch.html", prediction=predicted_role)

#     return render_template("JobSearch.html", prediction=None)

#-------------------------------------------------------------------------------------Menotor-bot----------------------------------------------------------------


app.secret_key = "1ccd2279c8c113b3d72605ee45c7be8f445133e614a65fda2a1401e2e4a79319"  # Required for session management

# Load chatbot model
model1 = tf.keras.models.load_model("chatbot_model.h5")

# Load words and classes
with open("words.pkl", "rb") as file:
    words = pickle.load(file)

with open("classes.pkl", "rb") as file:
    classes = pickle.load(file)

# Load intents
with open("intents.json", "r") as file:
    intents = json.load(file)

# Initialize lemmatizer
lemmatizer = WordNetLemmatizer()

# Configure Google Gemini API
GEMINI_API_KEY = "AIzaSyBS6FnIGfBxFyw3R58r36cX1p75zNVRexg"  # Replace with actual API key
genai.configure(api_key=GEMINI_API_KEY)

# Function to preprocess user input
def clean_up_sentence(sentence):
    sentence_words = nltk.word_tokenize(sentence)
    sentence_words = [lemmatizer.lemmatize(word.lower()) for word in sentence_words]
    return sentence_words

# Function to convert sentence into bag of words
def bow(sentence, words):
    sentence_words = clean_up_sentence(sentence)
    bag = [1 if word in sentence_words else 0 for word in words]
    return np.array(bag).reshape(1, -1)

# Function to predict intent
def predict_intent(sentence):
    bag = bow(sentence, words)
    res = model1.predict(bag)[0]
    max_index = np.argmax(res)

    print(f"Raw model output: {res}")  # Debugging line
    print(f"Predicted intent: {classes[max_index]}, Confidence: {res[max_index]}")  # Debugging line

    if res[max_index] > 0.5:  # Confidence threshold
        return classes[max_index]
    return "unknown"

# Function to get response based on intent
def get_response(intent):
    for intent_data in intents["intents"]:
        if intent_data["tag"] == intent:
            return random.choice(intent_data["responses"])
    return None

# Function to get response from Gemini API (with chat history)
def get_gemini_response(user_input):
    try:
        # Retrieve chat history from session
        chat_history = session.get("chat_history", [])

        # Prepare chat context
        chat_context = "\n".join(chat_history[-5:])  # Last 5 exchanges
        prompt = f"Previous conversation:\n{chat_context}\nUser: {user_input}\nAI:   but in max 40 words or 50 words "

        # Get response from Gemini
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)

        return response.text if response.text else "I'm not sure, could you clarify?"
    except Exception as e:
        return f"Error fetching response from Gemini: {str(e)}"

# Flask Route to get chatbot response
@app.route("/get_response", methods=["POST"])
def get_chatbot_response():
    data = request.get_json()
    user_message = data["message"]

    # Retrieve or initialize chat history
    chat_history = session.get("chat_history", [])

    # Predict intent
    intent = predict_intent(user_message)
    
    # Get predefined response
    response = get_response(intent)

    # If no response found, ask Gemini
    if not response or intent == "unknown":
        response = get_gemini_response(user_message)

    # Store chat history
    chat_history.append(f"User: {user_message}")
    chat_history.append(f"AI: {response}")
    session["chat_history"] = chat_history  # Update session

    return jsonify({"response": response})

# Route to clear chat history
@app.route("/clear_chat", methods=["POST"])
def clear_chat_history():
    session.pop("chat_history", None)
    return jsonify({"message": "Chat history cleared."})
#--------------------------------------------------------------------------------------------------------apply - job-search ----------------------------------------------------------------------
RAPIDAPI_HOST = "jsearch.p.rapidapi.com"
RAPIDAPI_KEY = "9a6fafe68cmshca09d2d8e9dbb59p14e3d5jsn2f64c416e74b"  # Replace with your actual API key

RAPIDAPI_URL = f"https://{RAPIDAPI_HOST}/search"


@app.route('/api/job_search', methods=['POST'])
def job_search():
    try:
        data = request.json
        search_query = data.get("input", "")
        print(search_query)
        if not search_query:
            return jsonify({"error": "Invalid input"}), 400

        # Define API request parameters
        query_params = {
            "query": search_query,
            "num_pages": 1  # Fetch 1 page of results (adjust if needed)
        }

        headers = {
            "X-RapidAPI-Key": RAPIDAPI_KEY,
            "X-RapidAPI-Host": RAPIDAPI_HOST
        }

        # Send request to RapidAPI
        response = requests.get(RAPIDAPI_URL, headers=headers, params=query_params)

        if response.status_code != 200:
            return jsonify({"error": "Failed to fetch job data"}), response.status_code

        job_data = response.json()

        # Extract relevant job details
        jobs = []
        for job in job_data.get("data", []):
            jobs.append({
                "title": job.get("job_title"),
                "company": job.get("employer_name"),
                "location": job.get("job_city", "N/A"),
                "apply_link": job.get("job_apply_link")
            })

        return jsonify({"jobs": jobs})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
#-------------------------------------------------------------------------------------interview Bot---------------------------------------------------






CORS(app)  # Enable CORS for frontend integration

genai.configure(api_key="AIzaSyBSwVAIWGui-QAxEzufmcpVH-RvWjrNmqc")
model = genai.GenerativeModel("gemini-1.5-flash")

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"pdf", "doc", "docx"}
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_resume_text(file_path):
    ext = file_path.rsplit(".", 1)[1].lower()

    try:
        if ext == "pdf":
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            return text if text else "No readable text found."

        elif ext == "docx":
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text if text else "No readable text found."

        elif ext == "doc":
            # Convert .doc to .docx using LibreOffice (Linux) or unoconv
            docx_path = file_path + "x"
            subprocess.run(["soffice", "--headless", "--convert-to", "docx", "--outdir", UPLOAD_FOLDER, file_path])
            if os.path.exists(docx_path):
                return extract_resume_text(docx_path)
            else:
                return "Failed to convert .doc file."

    except Exception as e:
        return f"Error extracting text: {str(e)}"

@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file format. Only PDF, DOC, DOCX allowed."}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    
    try:
        file.save(file_path)
        resume_text = extract_resume_text(file_path)
        os.remove(file_path)
        return jsonify({"resume_text": resume_text})
    except Exception as e:
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500


@app.route("/ask-questions", methods=["POST"])
def ask_questions():
    """Generates multiple interview questions based on the resume."""
    data = request.json
    resume_text = data.get("resume_text", "")
    num_questions = 5  # Number of questions to generate

    if not resume_text:
        return jsonify({"error": "No resume text provided"}), 400

    prompt = f"""
    You are an AI interviewer. Based on the following resume, generate {num_questions} unique and relevant interview questions.
    
    Resume:
    {resume_text}

    Rules:
    - Each question should be different from the others.
    - Focus on technical and behavioral aspects related to the resume.
    - Avoid generic questions like "Tell me about yourself."
    - Format each question as a separate line.

    Example Output:
    1. What challenges did you face while working on [specific project]?
    2. How do you optimize [specific skill] in your work?
    3. Explain a time when you handled a difficult situation using [specific skill].
    """

    try:
        response = model.generate_content(prompt)
        raw_questions = response.text.strip().split("\n")

        # Filter valid questions (removing empty lines, numbers, or unwanted text)
        questions = [q.strip() for q in raw_questions if q.strip() and not q.startswith("Example Output")]

        # Limit to the required number of questions
        questions = questions[:num_questions]

        if not questions:
            return jsonify({"error": "No valid questions generated."}), 500

    except Exception as e:
        return jsonify({"error": f"Error generating questions: {str(e)}"}), 500

    return jsonify({"questions": questions})


@app.route("/evaluate-answer", methods=["POST"])
def evaluate_answer():
    data = request.json
    resume_text = data.get("resume_text", "")
    question = data.get("question", "")
    answer = data.get("answer", "")

    prompt = f"""
    You are an interviewer evaluating a candidate's response.
    
    Question: {question}
    Answer: {answer}
    Resume:
    {resume_text}
    
    Task:
    1. Evaluate the correctness of the answer.
    2. Provide brief feedback.
    3. Rate the answer from 1 to 5.
    """

    try:
        response = model.generate_content(prompt)
        evaluation = response.text.strip()
    except Exception as e:
        return jsonify({"error": f"Error evaluating answer: {str(e)}"}), 500
    
    return jsonify({"evaluation": evaluation})

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=7860, debug=True)